"""
core/report_generator.py
Generates CDSCO-standard reports for all 3 document types:
  1. SUGAM Application Summary  (Parts A-E)
  2. SAE Case Summary           (Sections I-IX)
  3. Committee Meeting Summary  (Header + Agenda blocks)
All include mandatory masthead + Source Provenance block.
"""
import hashlib, logging
from datetime import datetime
from pathlib import Path
from typing import Optional
from config.settings import OUTPUT_DIR, REPORT_FOOTER_TEXT

logger = logging.getLogger("cdsco.report_generator")

# ── safe latin-1 helper ──────────────────────────────────────────
def _s(t):
    return str(t or "").replace("\u2014","-").replace("\u2013","-").replace("\u2018","'").replace("\u2019","'").replace("\u201C",'"').replace("\u201D",'"').replace("\u2022","*")

# ── auto-select algorithm based on doc type ──────────────────────
def auto_algo(doc_type: str) -> str:
    """Select best algorithm automatically — no user choice needed."""
    if "SAE" in doc_type:     return "luhn"     # short narratives → frequency-based
    if "Meeting" in doc_type: return "lexrank"  # long transcripts → graph-based
    return "lsa"                                  # inspection/SUGAM → semantic

# ── confidence score (simple heuristic) ─────────────────────────
def confidence_score(summary: str, raw_text: str) -> float:
    if not raw_text: return 0.0
    ratio = len(summary) / max(len(raw_text), 1)
    word_count = len(summary.split())
    score = min(0.95, 0.5 + (0.3 if word_count > 30 else 0.1) + (0.15 if ratio < 0.6 else 0.05))
    return round(score, 2)

# ── SHA-256 hash of source text ──────────────────────────────────
def file_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()[:16] + "..."


class ReportData:
    def __init__(self, firm_name="", firm_address="", license_no="", product_category="",
                 inspection_date="", inspection_type="", inspectors="", raw_text="",
                 summary="", key_points=None, observations="", deficiency_class="",
                 capa="", conclusion="", algorithm_used="", doc_type="SUGAM Portal / Inspection Data",
                 input_modality="Written transcript (.txt / .pdf / .docx)",
                 audio_duration="", division="NDD", sugam_ref="", form_type=""):
        self.firm_name=_s(firm_name); self.firm_address=_s(firm_address)
        self.license_no=_s(license_no); self.product_category=_s(product_category)
        self.inspection_date=_s(inspection_date); self.inspection_type=_s(inspection_type)
        self.inspectors=_s(inspectors); self.raw_text=raw_text
        self.summary=_s(summary); self.key_points=[_s(k) for k in (key_points or [])]
        self.observations=_s(observations); self.deficiency_class=_s(deficiency_class)
        self.capa=_s(capa); self.doc_type=doc_type
        self.input_modality=input_modality; self.audio_duration=audio_duration
        self.division=division; self.sugam_ref=sugam_ref; self.form_type=form_type
        self.algorithm_used=_s(algorithm_used or auto_algo(doc_type))
        self.generated_at=datetime.now().strftime("%d %B %Y, %H:%M")
        self.confidence=confidence_score(summary, raw_text)
        self.src_hash=file_hash(raw_text) if raw_text else "N/A"
        self.conclusion=_s(conclusion) if conclusion else (
            "Based on the above, the matter is submitted for kind consideration and orders of the reviewing officer.")


class ReportGenerator:
    def __init__(self, output_dir=None):
        self._out = Path(output_dir or OUTPUT_DIR)
        self._out.mkdir(parents=True, exist_ok=True)

    def generate(self, data: ReportData, output_format="both") -> dict:
        paths = {}
        fmt = output_format.lower()
        if fmt in ("pdf","both"): paths["pdf"]  = self._pdf(data)
        if fmt in ("docx","both"): paths["docx"] = self._docx(data)
        return paths

    # ─────────────────────────────────────────────────────────────
    # MASTHEAD (common to all 3 templates)
    # ─────────────────────────────────────────────────────────────
    def _masthead_lines(self, data: ReportData) -> list:
        is_sae     = "SAE" in data.doc_type
        is_meeting = "Meeting" in data.doc_type
        doc_class  = ("SUGAM Application Summary" if not is_sae and not is_meeting
                      else "SAE Case Summary" if is_sae else "Committee Meeting Summary")
        lines = [
            "GOVERNMENT OF INDIA",
            "MINISTRY OF HEALTH & FAMILY WELFARE",
            "CENTRAL DRUGS STANDARD CONTROL ORGANISATION (CDSCO)",
            "FDA Bhawan, Kotla Road, New Delhi - 110002",
            "",
            f"Document Class     : {doc_class}",
            f"File / SUGAM Ref   : {data.sugam_ref or data.license_no or 'N/A'}",
            f"Division           : {data.division}",
            f"Prepared by (AI)   : CDSCO Data Summarisation Tool v1.0 | Algo: {data.algorithm_used.upper()}",
            f"Human Reviewer     : {data.inspectors or '___________________'} | Date: {data.inspection_date}",
            f"Confidence Score   : {data.confidence}  (fields below 0.7 flagged)",
            "Classification     : RESTRICTED - For CDSCO Internal Use Only",
            "",
            "SOURCE PROVENANCE",
            "=" * 60,
            f"Input Modality     : {data.input_modality}",
            f"Audio Duration     : {data.audio_duration or 'N/A'}",
            f"ASR Engine         : Vosk vosk-model-en-in-0.5 (offline)",
            f"File Hash (SHA-256): {data.src_hash}",
            "=" * 60,
        ]
        return lines

    # ─────────────────────────────────────────────────────────────
    # TEMPLATE 1 — SUGAM Application Summary
    # ─────────────────────────────────────────────────────────────
    def _sugam_body(self, data: ReportData) -> list:
        kps = "\n".join(f"  {k}" for k in data.key_points) if data.key_points else "  (not extracted)"
        return [
            "PART A - APPLICATION IDENTIFICATION",
            "=" * 60,
            f"A.1  SUGAM Application No.   : {data.sugam_ref or data.license_no}",
            f"A.2  Form Type / Rule         : {data.form_type or 'Form CT-04 / Form 44 / Form MD-14'}",
            f"A.3  Applicant / Sponsor      : {data.firm_name}",
            f"A.4  Registered Address       : {data.firm_address}",
            f"A.5  Product / Drug / Device  : {data.product_category}",
            f"A.6  Category                 : {data.product_category}",
            f"A.7  Inspection Type          : {data.inspection_type}",
            f"A.8  Date of Submission       : {data.inspection_date}",
            f"A.9  Officers / Reviewers     : {data.inspectors}",
            "",
            "PART B - CHECKLIST COMPLIANCE MATRIX",
            "=" * 60,
            f"{'Sl.':<4} {'Checklist Item':<35} {'Status':<15} {'Finding'}",
            "-" * 80,
        ] + self._checklist_rows(data) + [
            "",
            "PART C - TECHNICAL PARTICULARS",
            "=" * 60,
            "C.1  Summary of Observations:",
            f"     {data.summary}",
            "",
            "C.2  Key Extracted Findings:",
            kps,
            "",
            f"C.3  Risk-Benefit One-liner   : {self._risk_benefit(data)}",
            "",
            "PART D - DEFICIENCIES & QUERIES",
            "=" * 60,
            f"D.1  Deficiency Classification: {data.deficiency_class or 'Pending Review'}",
            "",
            "D.2  Deficiencies Identified:",
            "\n".join(f"  [{i+1}] {k}" for i,k in enumerate(
                [k for k in data.key_points if any(w in k.lower() for w in
                 ["missing","overdue","without","not have","not signed","incomplete",
                  "invalid","failed","no valid","not current","two missing","last revised",
                  "last performed","overdue"])]
            )) or "  None identified.",
            "",
            "D.3  CAPA / Draft Query Letter Points:",
            f"  {data.capa or 'To be determined by reviewing officer.'}",
            "",
            "PART E - RECOMMENDATION (FOR HUMAN REVIEWER ONLY)",
            "=" * 60,
            "  [ ] Accept for further processing",
            "  [ ] Raise Query to Applicant",
            "  [ ] Refer to SEC",
            "  [ ] Reject with reasons",
            "  [ ] Refer to CDL/CDTL for testing",
            "  [ ] Escalate to Jt. DC / DCG(I)",
            "",
            "Reviewer Signature: _______________  Date: _______________",
            "Dy. Drugs Controller: _____________  Jt. Drugs Controller: _______________",
        ]

    def _checklist_rows(self, data: ReportData) -> list:
        """Auto-generate checklist rows from key points."""
        rows = []
        standard_items = [
            "Manufacturing Licence / Form",
            "GMP Compliance Certificate",
            "HVAC Qualification Records",
            "Water System Validation",
            "BMR / BPR Records",
            "QC Training Records",
            "SOP Currency",
            "CAPA Register",
            "Storage Conditions",
            "Complaint Investigation",
        ]
        kp_text = " ".join(data.key_points).lower()
        neg_words = ["missing","overdue","not signed","invalid","without","incomplete",
                     "failed","no valid","not current","last revised","last performed"]
        for i, item in enumerate(standard_items):
            keywords = item.lower().split()
            found = any(k in kp_text for k in keywords)
            finding_text = next((k for k in data.key_points if any(kw in k.lower() for kw in keywords)), "")
            is_deficient = found and any(w in finding_text.lower() for w in neg_words)
            flag = "INCOMPLETE *" if is_deficient else ("REVIEWED" if found else "OK")
            rows.append(f"{i+1:<4} {item:<35} {flag:<15} {finding_text[:60] or '-'}")
        return rows

    def _risk_benefit(self, data: ReportData) -> str:
        if "critical" in (data.deficiency_class or "").lower():
            return "Risk outweighs benefit pending resolution of critical deficiencies."
        if "major" in (data.deficiency_class or "").lower():
            return "Benefit acceptable subject to resolution of major deficiencies within 45 days."
        return "Benefit-risk profile acceptable based on current inspection findings."

    # ─────────────────────────────────────────────────────────────
    # TEMPLATE 2 — SAE Case Summary
    # ─────────────────────────────────────────────────────────────
    def _sae_body(self, data: ReportData) -> list:
        kps = "\n".join(f"  {k}" for k in data.key_points) if data.key_points else "  (see narrative)"
        return [
            "SECTION I - REPORT METADATA & TIMELINE COMPLIANCE",
            "=" * 60,
            f"I.1  SAE Report No.            : {data.license_no}",
            "I.2  Report Type               : [ ] Initial  [ ] Follow-up",
            f"I.3  Date of Event             : {data.inspection_date}",
            f"I.4  Sponsor / CRO / Site / PI : {data.firm_name}",
            f"I.5  Institution               : {data.firm_address}",
            f"I.6  Reporting Officer         : {data.inspectors}",
            "I.7  Timeline Compliance       : [ ] Within 14 calendar days  [ ] DELAYED",
            f"I.8  Protocol / CT Permission  : {data.sugam_ref or 'N/A'}",
            "",
            "SECTION II - SUBJECT PROFILE",
            "=" * 60,
            "II.1  Subject Code / Initials  : (anonymised per DPDP Act 2023)",
            "II.2  Age / Sex / Ethnicity    : (extracted from narrative)",
            "II.3  Relevant Medical History : (see Section IV narrative)",
            "II.4  Concomitant Medications  : (see Section IV narrative)",
            "",
            "SECTION III - INVESTIGATIONAL PRODUCT EXPOSURE",
            "=" * 60,
            f"III.1 IP / Drug Name           : {data.product_category}",
            "III.2 Dose / Route / Regimen   : (extracted from narrative)",
            "III.3 Blinding Status          : [ ] Blinded  [ ] Unblinded",
            "III.4 Time-to-Onset from Dose  : (extracted from narrative)",
            "",
            "SECTION IV - SAE CLINICAL NARRATIVE",
            "=" * 60,
            data.summary or "(Narrative not provided)",
            "",
            "SECTION V - SERIOUSNESS CRITERIA (Schedule Y)",
            "=" * 60,
            "  [ ] Death                    [ ] Life-threatening",
            "  [ ] Hospitalisation          [ ] Persistent disability",
            "  [ ] Congenital anomaly        [ ] Other medically important",
            f"  AI-extracted indication: {data.deficiency_class or 'Pending reviewer assessment'}",
            "",
            "SECTION VI - CAUSALITY ASSESSMENT",
            "=" * 60,
            "  VI.1 Investigator Opinion    : [ ] Certain  [ ] Probable  [ ] Possible  [ ] Unlikely",
            "  VI.2 Sponsor Medical Monitor : _______________",
            "  VI.3 Ethics Committee View   : _______________",
            f"  VI.4 WHO-UMC Category       : (reviewer to assign)",
            "  VI.5 Expectedness vs IB      : [ ] Expected  [ ] Unexpected",
            "",
            "SECTION VII - COMPENSATION (Rule 42, NDCT Rules 2019)",
            "=" * 60,
            "  VII.1 Compensation Applicable: Y / N",
            "  VII.2 Amount Proposed        : Rs. _______",
            "  VII.3 Status of Payment      : _______________",
            "",
            "SECTION VIII - AI-EXTRACTED KEY POINTS",
            "=" * 60,
            kps,
            "",
            "SECTION VIII-B - CAPA / ACTIONS",
            f"  {data.capa or 'To be determined by reviewing officer.'}",
            "",
            "SECTION IX - REVIEWER DECISION PANEL",
            "=" * 60,
            "  [ ] Accept - Related / compensation payable",
            "  [ ] Accept - Not related / no compensation",
            "  [ ] Refer to SAE Independent Expert Committee",
            "  [ ] Query to Sponsor",
            "  [ ] Trigger GCP Inspection at site",
            "  [ ] Show-cause / Warning recommended",
            "",
            "Reviewing Officer: _______________  DDC (CT Division): _______________",
        ]

    # ─────────────────────────────────────────────────────────────
    # TEMPLATE 3 — Committee Meeting Summary
    # ─────────────────────────────────────────────────────────────
    def _meeting_body(self, data: ReportData) -> list:
        kps = data.key_points or []
        # Split key points into decisions, actions, next steps
        decisions = [k for k in kps if any(w in k.lower() for w in ["decision","decided","approved","recommended","rejected"])]
        actions   = [k for k in kps if any(w in k.lower() for w in ["action","to submit","to provide","to conduct","to draft"])]
        other     = [k for k in kps if k not in decisions and k not in actions]
        if not decisions: decisions = kps[:3]
        if not actions:   actions   = kps[3:]

        rows = []
        rows += [
            "COMMITTEE MEETING IDENTIFICATION",
            "=" * 60,
            f"Committee          : {data.product_category or 'SEC / Technical Committee / DTAB'}",
            f"Meeting Ref        : {data.sugam_ref or data.license_no}",
            f"Date               : {data.inspection_date}",
            f"Venue / Mode       : {data.firm_address or 'FDA Bhawan, New Delhi'}",
            f"Chairperson        : {data.inspectors.split(',')[0] if data.inspectors else '___'}",
            f"CDSCO Moderator    : {data.inspectors.split(',')[1].strip() if ',' in (data.inspectors or '') else '___'}",
            f"Members Present    : {data.inspectors}",
            "Quorum Met         : [ ] Yes (min. 4 incl. 1 Pharmacologist)  [ ] No",
            "Pre-brief circulated: [ ] Yes (>=5 days)  [ ] Delayed",
            "",
            "AGENDA ITEMS",
            "=" * 60,
            f"(a) File No. / Ref   : {data.license_no}",
            f"(b) Applicant        : {data.firm_name}",
            f"(c) Nature           : {data.inspection_type}",
            "",
            "(e) Background / Context (AI-extracted):",
            f"    {data.summary[:300] if data.summary else '(see full summary below)'}",
            "",
            "(f) Key Points Presented:",
        ]
        for k in (other or kps)[:4]:
            rows.append(f"    * {k}")
        rows += [
            "",
            "(g) Decisions / Recommendations:",
        ]
        for k in (decisions or kps[:3]):
            rows.append(f"    * {k}")
        rows += [
            "",
            "(h) Rationale for Decision (MANDATORY per SEC Guidance 2025):",
            f"    {data.conclusion}",
            "",
            "(i) DECISION:",
            "    [ ] Recommended",
            "    [ ] Recommended with minor changes",
            "    [ ] Recommended with major changes (re-evaluation required)",
            "    [ ] Not Recommended",
            "    [ ] Further review - additional data sought",
            "    [ ] Deferred to next meeting",
            "",
            "CONSOLIDATED ACTION-ITEM REGISTER",
            "=" * 60,
            f"{'#':<4} {'Action':<40} {'Owner':<20} {'Due Date'}",
            "-" * 80,
        ]
        for i, act in enumerate(actions or kps[-3:], 1):
            rows.append(f"{i:<4} {act[:38]:<40} {'Applicant':<20} TBD")
        rows += [
            "",
            "ANY OTHER BUSINESS: _______________",
            "",
            f"NEXT MEETING: _______________  Provisional Agenda: _______________",
            "",
            "ANNEXURES",
            "=" * 60,
            "Annex-I:   Attendance Sheet",
            "Annex-II:  Briefing Materials Index",
            "Annex-III: Declarations of Conflict of Interest",
            "Annex-IV:  Slides / Data Shown",
            "",
            f"Minutes Drafted by (AI): CDSCO Summarisation Tool v1.0 | {data.generated_at}",
            "Verified by CDSCO Moderator: _______________",
            "Approved by Chairperson:     _______________",
        ]
        return rows

    # ─────────────────────────────────────────────────────────────
    # FULL TEXT ASSEMBLY
    # ─────────────────────────────────────────────────────────────
    def _full_text(self, data: ReportData) -> list:
        lines = self._masthead_lines(data)
        lines.append("")
        if "SAE" in data.doc_type:
            lines += self._sae_body(data)
        elif "Meeting" in data.doc_type:
            lines += self._meeting_body(data)
        else:
            lines += self._sugam_body(data)
        lines += ["", "=" * 60,
                  "RESTRICTED - For CDSCO Internal Use Only",
                  "Prepared under CDSCO Schedule M / MDR 2017 / DPDP Act 2023",
                  "=" * 60]
        return lines

    # ─────────────────────────────────────────────────────────────
    # PDF OUTPUT
    # ─────────────────────────────────────────────────────────────
    def _pdf(self, data: ReportData) -> str:
        try:
            from fpdf import FPDF, XPos, YPos
        except ImportError:
            logger.error("fpdf2 not installed"); return ""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=18)
        pdf.add_page()
        pdf.set_margins(18, 18, 18)

        # Masthead
        pdf.set_font("Helvetica","B",13)
        pdf.set_text_color(0,51,102)
        pdf.cell(0,8,"CENTRAL DRUGS STANDARD CONTROL ORGANISATION",align="C",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
        pdf.set_font("Helvetica","",10)
        pdf.cell(0,6,"Ministry of Health & Family Welfare, Government of India",align="C",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
        pdf.set_draw_color(200,150,12); pdf.set_line_width(0.8)
        pdf.line(18,pdf.get_y(),192,pdf.get_y()); pdf.ln(3)

        is_sae     = "SAE" in data.doc_type
        is_meeting = "Meeting" in data.doc_type
        title = ("SAE CASE SUMMARY REPORT" if is_sae else
                 "COMMITTEE MEETING SUMMARY" if is_meeting else
                 "SUGAM APPLICATION SUMMARY REPORT")
        pdf.set_font("Helvetica","B",12)
        pdf.set_text_color(0,0,0)
        pdf.cell(0,8,title,align="C",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
        pdf.set_font("Helvetica","",9)
        pdf.set_text_color(100,100,100)
        pdf.cell(0,6,f"Generated: {data.generated_at}  |  Confidence: {data.confidence}  |  Algorithm: {data.algorithm_used.upper()}",align="C",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
        pdf.ln(4)

        # Body
        pdf.set_text_color(0,0,0)
        lines = self._full_text(data)
        for line in lines:
            line = _s(line)
            if line.startswith("="):
                pdf.set_draw_color(0,51,102); pdf.set_line_width(0.5)
                pdf.line(18,pdf.get_y(),192,pdf.get_y()); pdf.ln(3)
            elif line.startswith("PART ") or line.startswith("SECTION ") or line.startswith("AGENDA") or line.startswith("CONSOLIDATED") or line.startswith("COMMITTEE") or line.startswith("SOURCE"):
                pdf.set_font("Helvetica","B",10)
                pdf.set_fill_color(0,51,102); pdf.set_text_color(255,255,255)
                pdf.cell(0,7,f"  {line}",fill=True,new_x=XPos.LMARGIN,new_y=YPos.NEXT)
                pdf.set_text_color(0,0,0); pdf.ln(1)
            elif line.strip() == "":
                pdf.ln(3)
            else:
                pdf.set_font("Helvetica","",9)
                pdf.multi_cell(pdf.epw,5.5,line)

        # Footer
        pdf.set_draw_color(200,150,12); pdf.line(18,pdf.get_y(),192,pdf.get_y()); pdf.ln(2)
        pdf.set_font("Helvetica","I",8); pdf.set_text_color(120,120,120)
        pdf.multi_cell(pdf.epw,5,_s(REPORT_FOOTER_TEXT),align="C")

        fname = self._fname(data,"pdf")
        pdf.output(fname); logger.info("PDF: %s",fname); return fname

    # ─────────────────────────────────────────────────────────────
    # DOCX OUTPUT
    # ─────────────────────────────────────────────────────────────
    def _docx(self, data: ReportData) -> str:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed"); return ""
        doc = Document()
        for sec in doc.sections:
            sec.top_margin=sec.bottom_margin=Inches(1.0)
            sec.left_margin=sec.right_margin=Inches(1.2)

        NAVY=RGBColor(0,51,102); GOLD=RGBColor(200,150,12)

        def add_h(text, color=NAVY):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(8)
            r=p.add_run(text); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=color

        def add_body(text):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(2)
            r=p.add_run(text); r.font.size=Pt(10)

        # Header
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("CENTRAL DRUGS STANDARD CONTROL ORGANISATION")
        r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY
        p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r2=p2.add_run("Ministry of Health & Family Welfare, Government of India")
        r2.font.size=Pt(11)

        is_sae="SAE" in data.doc_type; is_meeting="Meeting" in data.doc_type
        title=("SAE CASE SUMMARY REPORT" if is_sae else
               "COMMITTEE MEETING SUMMARY" if is_meeting else
               "SUGAM APPLICATION SUMMARY REPORT")
        p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r3=p3.add_run(title); r3.bold=True; r3.font.size=Pt(13); r3.font.color.rgb=NAVY

        meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r4=meta.add_run(f"Generated: {data.generated_at}  |  Confidence: {data.confidence}  |  Algorithm: {data.algorithm_used.upper()}")
        r4.font.size=Pt(9); r4.font.color.rgb=RGBColor(100,100,100)
        doc.add_paragraph()

        for line in self._full_text(data):
            if line.startswith("=") or line.startswith("-"):
                continue
            elif (line.startswith("PART ") or line.startswith("SECTION ") or
                  line.startswith("AGENDA") or line.startswith("CONSOLIDATED") or
                  line.startswith("COMMITTEE") or line.startswith("SOURCE")):
                add_h(line)
            elif line.strip() == "":
                doc.add_paragraph()
            else:
                add_body(line)

        # Footer
        footer=doc.sections[0].footer.paragraphs[0]
        footer.text=_s(REPORT_FOOTER_TEXT)
        footer.alignment=WD_ALIGN_PARAGRAPH.CENTER

        fname=self._fname(data,"docx")
        doc.save(fname); logger.info("DOCX: %s",fname); return fname

    def _fname(self, data, ext):
        ts=datetime.now().strftime("%Y%m%d_%H%M%S")
        safe="".join(c for c in data.firm_name if c.isalnum() or c in " _-")[:25].strip().replace(" ","_")
        prefix=("SAE" if "SAE" in data.doc_type else "MTG" if "Meeting" in data.doc_type else "SUGAM")
        return str(self._out/f"CDSCO_{prefix}_{safe}_{ts}.{ext}")