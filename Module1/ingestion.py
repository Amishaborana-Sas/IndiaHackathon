"""
Module 1 — Ingestion Layer
==========================
Universal input handler. Every heavy dependency is imported lazily so a
missing one only breaks the specific format that needs it — everything
else keeps working.

OCR backend: EasyOCR (default) or PaddleOCR — configurable via
ocr_config.py or the OCR_ENGINE environment variable.
Tesseract is NOT used anywhere in this module.

Return format (always the same)
-------------------------------
{
    "source":      str,
    "source_type": str,
    "is_scanned":  bool,
    "text":        str,
    "structured":  list[dict] | None,
    "metadata":    dict,
}
"""

from __future__ import annotations

import csv
import io
import json
import logging
import os
import re
import sys
import time
from pathlib import Path
from typing import Any, Union

import ocr_config

PathLike = Union[str, os.PathLike]
logger = logging.getLogger(__name__)

# ========================================================================== #
#  SHARED OCR ENGINE — uses the user's ocr_module/ocr_engine.py             #
# ========================================================================== #
_shared_ocr_engine = None


def _get_shared_ocr():
    """Get the shared OCR engine (lazy loaded)."""
    global _shared_ocr_engine
    if _shared_ocr_engine is None:
        try:
            _ocr_dir = str(Path(__file__).resolve().parent.parent / "ocr_module")
            if _ocr_dir not in sys.path:
                sys.path.insert(0, _ocr_dir)
            from ocr_engine import OCREngine
            _shared_ocr_engine = OCREngine(languages=["en"], gpu=False, dpi=200, enhance=True)
            logger.info("Shared OCR Engine loaded (EasyOCR + PyMuPDF)")
        except Exception as e:
            logger.warning("Could not load shared OCR engine: %s", e)
    return _shared_ocr_engine


def _shared_ocr_extract(data: bytes, filename: str) -> str | None:
    """Extract text using the shared OCR engine. Returns None if unavailable."""
    engine = _get_shared_ocr()
    if engine is None:
        return None
    try:
        result = engine.run(data, filename=filename)
        text = result.get("full_text", "")
        if text and len(text.strip()) > 10:
            logger.info("Shared OCR: %d chars from %s (%.1fs, type=%s)",
                       len(text), filename, result.get("elapsed_sec", 0),
                       result.get("source_type", "?"))
            return text
    except Exception as e:
        logger.warning("Shared OCR failed for %s: %s", filename, e)
    return None


class IngestionError(Exception):
    """Raised when a file can't be ingested (with a friendly message)."""


def _missing(pkg: str, fmt: str, install: str) -> IngestionError:
    return IngestionError(
        f"Cannot ingest {fmt} — '{pkg}' is not installed.\n"
        f"  Install it with:  {install}"
    )


# ========================================================================== #
#  OCR ENGINE — Singleton wrapper for EasyOCR / PaddleOCR                    #
# ========================================================================== #
_ocr_reader = None          # module-level singleton — loaded once
_ocr_engine_name: str = ""  # tracks which engine is currently loaded


def _get_ocr_reader(engine_name: str | None = None):
    """
    Return the cached OCR reader. Creates it on first call.
    Models are large so we NEVER re-create unless the engine name changes.
    """
    global _ocr_reader, _ocr_engine_name
    name = (engine_name or ocr_config.OCR_ENGINE).lower().strip()

    if _ocr_reader is not None and _ocr_engine_name == name:
        return _ocr_reader  # cache hit

    logger.info("Loading OCR engine: %s (gpu=%s) …", name, ocr_config.USE_GPU)
    t0 = time.perf_counter()

    if name == "easyocr":
        try:
            import easyocr
        except ImportError:
            raise _missing("easyocr", "images / scanned PDFs",
                           "pip install easyocr")
        _ocr_reader = easyocr.Reader(
            ocr_config.OCR_LANGUAGES,
            gpu=ocr_config.USE_GPU,
            verbose=False,
        )
    elif name in ("paddleocr", "paddle"):
        try:
            from paddleocr import PaddleOCR
        except ImportError:
            raise _missing("paddleocr", "images / scanned PDFs",
                           "pip install paddlepaddle paddleocr")
        lang = ocr_config.OCR_LANGUAGES[0] if ocr_config.OCR_LANGUAGES else "en"
        _ocr_reader = PaddleOCR(
            use_angle_cls=True,
            lang=lang,
            use_gpu=ocr_config.USE_GPU,
            show_log=False,
        )
    else:
        raise IngestionError(f"Unknown OCR engine: {name}. Use 'easyocr' or 'paddleocr'.")

    _ocr_engine_name = name
    logger.info("OCR engine %s loaded in %.1fs", name, time.perf_counter() - t0)
    return _ocr_reader


def get_ocr_engine_name() -> str:
    """Return which OCR engine is configured (for health endpoint)."""
    return _ocr_engine_name or ocr_config.OCR_ENGINE


# ========================================================================== #
#  IMAGE PREPROCESSING — OpenCV pipeline                                     #
# ========================================================================== #
def _preprocess_image(img):
    """
    Full preprocessing pipeline: resize → grayscale → denoise →
    adaptive threshold → deskew.

    Each step is toggleable via ocr_config. Returns a cleaned image
    ready for OCR. Never raises — returns the input on failure.
    """
    import cv2
    import numpy as np

    if img is None:
        raise IngestionError("Empty image passed to preprocessor")

    # ── Resize if too large ──
    h, w = img.shape[:2]
    max_dim = ocr_config.MAX_IMAGE_DIMENSION
    if max(h, w) > max_dim:
        scale = max_dim / max(h, w)
        img = cv2.resize(img, (int(w * scale), int(h * scale)),
                         interpolation=cv2.INTER_AREA)

    # ── Grayscale ──
    if ocr_config.ENABLE_GRAYSCALE and len(img.shape) == 3:
        img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # ── Denoise ──
    if ocr_config.ENABLE_DENOISE:
        if len(img.shape) == 3:
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        img = cv2.fastNlMeansDenoising(
            img, None,
            h=ocr_config.DENOISE_STRENGTH,
            templateWindowSize=7,
            searchWindowSize=21,
        )

    # ── Adaptive threshold ──
    if ocr_config.ENABLE_THRESHOLD:
        if len(img.shape) == 3:
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        img = cv2.adaptiveThreshold(
            img, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            blockSize=ocr_config.THRESHOLD_BLOCK_SIZE,
            C=ocr_config.THRESHOLD_CONSTANT,
        )

    # ── Deskew ──
    if ocr_config.ENABLE_DESKEW:
        try:
            if len(img.shape) == 3:
                img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            coords = np.column_stack(np.where(img < 128))
            if coords.shape[0] > 50:
                angle = cv2.minAreaRect(coords)[-1]
                if angle < -45:
                    angle = 90 + angle
                if 0.3 < abs(angle) < 15:
                    (rh, rw) = img.shape[:2]
                    M = cv2.getRotationMatrix2D((rw // 2, rh // 2), angle, 1.0)
                    img = cv2.warpAffine(
                        img, M, (rw, rh),
                        flags=cv2.INTER_CUBIC,
                        borderMode=cv2.BORDER_REPLICATE,
                    )
                    logger.debug("Deskewed by %.2f°", angle)
        except Exception:
            logger.debug("Deskew failed, continuing with original")

    return img


# ========================================================================== #
#  HANDWRITTEN NOTES DETECTION — OpenCV heuristic                            #
# ========================================================================== #
def _detect_handwritten_regions(img) -> list[dict]:
    """
    Detect handwritten regions using OpenCV heuristics:
    Canny edge detection + contour analysis + ink density + stroke irregularity.
    Returns list of {"bbox": (x,y,w,h), "ink_density": float, "confidence": float}.
    """
    try:
        import cv2
        import numpy as np
    except ImportError:
        return []

    try:
        # Convert to grayscale if needed
        if len(img.shape) == 3:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        else:
            gray = img.copy()

        h, w = gray.shape[:2]
        if h < 50 or w < 50:
            return []

        # Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)

        # Canny edge detection
        edges = cv2.Canny(blurred, 50, 150)

        # Dilate edges to merge nearby strokes
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (7, 7))
        dilated = cv2.dilate(edges, kernel, iterations=3)

        # Find contours
        contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL,
                                        cv2.CHAIN_APPROX_SIMPLE)

        regions = []
        for cnt in contours:
            area = cv2.contourArea(cnt)
            if area < 800:
                continue

            x, y, rw, rh = cv2.boundingRect(cnt)

            # Skip very thin or very wide regions (likely printed lines/borders)
            aspect = rw / max(rh, 1)
            if aspect > 15 or aspect < 0.05:
                continue

            # Skip regions that are too large (likely the whole page)
            if rw > w * 0.9 and rh > h * 0.9:
                continue

            # Extract ROI and compute ink density
            roi = gray[y:y+rh, x:x+rw]
            dark_pixels = np.sum(roi < 128)
            total_pixels = roi.size
            ink_density = dark_pixels / max(total_pixels, 1)

            # Ink density between 0.03 and 0.55 suggests handwriting
            if ink_density < 0.03 or ink_density > 0.55:
                continue

            # Stroke irregularity via HoughLinesP
            roi_edges = cv2.Canny(roi, 50, 150)
            lines = cv2.HoughLinesP(roi_edges, 1, np.pi / 180,
                                     threshold=20, minLineLength=10, maxLineGap=5)
            angle_variance = 0.0
            if lines is not None and len(lines) > 3:
                angles = []
                for line in lines:
                    x1, y1, x2, y2 = line[0]
                    angle = np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi
                    angles.append(angle)
                angle_variance = np.var(angles)

            # High angle variance = irregular strokes = likely handwritten
            # Low angle variance = uniform text = likely printed
            confidence = 0.0
            if angle_variance > 200:
                confidence = min(0.95, 0.5 + angle_variance / 2000)
            elif angle_variance > 50:
                confidence = 0.3 + angle_variance / 1000

            if confidence >= 0.3:
                regions.append({
                    "bbox": (x, y, rw, rh),
                    "ink_density": round(ink_density, 4),
                    "confidence": round(confidence, 3),
                })

        return regions

    except Exception:
        logger.debug("Handwriting detection failed, skipping")
        return []


def _anonymize_handwritten_regions(img, regions: list[dict]):
    """Black out detected handwritten regions on the image.
    Returns (modified_image, count_of_regions_blanked)."""
    import cv2
    modified = img.copy()
    count = 0
    for region in regions:
        x, y, w, h = region["bbox"]
        cv2.rectangle(modified, (x, y), (x + w, y + h), (0, 0, 0), -1)
        count += 1
    return modified, count


# ========================================================================== #
#  SCANNED PDF DETECTION — multi-signal heuristic                            #
# ========================================================================== #
def _is_scanned_pdf(page_texts: list[str], page_count: int) -> bool:
    """
    Multi-signal heuristic to determine if a PDF is scanned.
    Checks text-per-page ratio, character variety, and word-like token ratio.
    """
    if page_count == 0:
        return True

    total_text = "\n".join(page_texts).strip()
    total_chars = len(total_text)

    # Signal 1: Very little text per page
    chars_per_page = total_chars / page_count
    if chars_per_page < 50:
        return True

    # Signal 2: Low character variety (garbage text layer)
    non_ws = re.sub(r"\s+", "", total_text)
    if non_ws:
        unique_chars = len(set(non_ws))
        if unique_chars < 10:
            return True

    # Signal 3: Low word-like token ratio
    tokens = total_text.split()
    if tokens:
        word_like = sum(1 for t in tokens if re.match(r"[a-zA-Z]{2,}", t))
        ratio = word_like / len(tokens)
        if ratio < 0.20:
            return True

    return False


# ========================================================================== #
#  TEXT POST-PROCESSING — clean OCR output for NER                           #
# ========================================================================== #
def _postprocess_ocr_text(raw: str) -> str:
    """
    Clean raw OCR text for better downstream NER/PII detection.

    Fixes: OCR char errors, broken lines, spacing, Indian identifier
    formatting (Aadhaar, PAN, phone).
    """
    if not raw or not raw.strip():
        return ""

    text = raw

    # ── Common OCR substitution errors ──
    for old, new in {
        "|": "I", "¢": "c", "©": "C",
        "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-",
        "\u00a0": " ",
    }.items():
        text = text.replace(old, new)

    # ── Collapse 3+ newlines into paragraph break ──
    text = re.sub(r"\n{3,}", "\n\n", text)

    # ── Join mid-sentence line breaks ──
    lines = text.split("\n")
    merged: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            merged.append("")
            continue
        if merged and merged[-1] and not re.search(r"[.!?:;]\s*$", merged[-1]):
            merged[-1] = merged[-1].rstrip() + " " + stripped
        else:
            merged.append(stripped)
    text = "\n".join(merged)

    # ── Rejoin hyphenated words across lines ──
    text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)

    # ── Fix spacing ──
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    text = re.sub(r"([.,;:!?])([A-Za-z])", r"\1 \2", text)

    # ── Normalise Indian identifiers ──
    # Aadhaar: XXXX XXXX XXXX
    text = re.sub(r"\b(\d{4})[.\-\s]?(\d{4})[.\-\s]?(\d{4})\b", r"\1 \2 \3", text)
    # PAN: ABCDE1234F (remove stray spaces)
    text = re.sub(
        r"\b([A-Z]{3})\s?([A-Z])\s?([A-Z])\s?(\d{4})\s?([A-Z])\b",
        r"\1\2\3\4\5", text,
    )
    # Phone: +91 XXXXX XXXXX
    text = re.sub(
        r"(\+91|0)\s*[-.]?\s*(\d{5})\s*[-.]?\s*(\d{5})",
        r"\1 \2 \3", text,
    )

    # ── Strip non-printable (keep ASCII + Devanagari) ──
    text = re.sub(r"[^\S\n]+", " ", text)
    text = re.sub(r"[^\x20-\x7E\n\u0900-\u097F]", "", text)

    # ── Final trim ──
    text = "\n".join(line.strip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ========================================================================== #
#  OCR EXTRACTION — run engine on preprocessed image                         #
# ========================================================================== #
def _ocr_image(img, engine_name: str | None = None) -> tuple[str, int]:
    """
    Preprocess image, detect/redact handwriting, run OCR, postprocess text.
    Returns (cleaned_text, handwritten_regions_count).
    Never crashes — returns ("", 0) on failure.
    """
    try:
        import cv2
        import numpy as np
    except ImportError:
        raise _missing("opencv-python", "images / scanned PDFs",
                       "pip install opencv-python-headless")

    try:
        processed = _preprocess_image(img)
    except IngestionError:
        raise
    except Exception:
        logger.warning("Preprocessing failed, using raw image for OCR")
        processed = img

    # Detect and redact handwritten regions before OCR
    hw_count = 0
    hw_regions = _detect_handwritten_regions(processed)
    if hw_regions:
        processed, hw_count = _anonymize_handwritten_regions(processed, hw_regions)
        logger.info("Handwriting: %d region(s) detected and redacted", hw_count)

    # Get the cached OCR reader
    reader = _get_ocr_reader(engine_name)
    name = _ocr_engine_name

    raw_text = ""
    try:
        if name == "easyocr":
            results = reader.readtext(processed, detail=0, paragraph=True)
            raw_text = "\n".join(results)
        elif name in ("paddleocr", "paddle"):
            results = reader.ocr(processed, cls=True)
            if results and results[0]:
                lines = []
                for line in results[0]:
                    if line and len(line) >= 2:
                        txt = line[1][0] if isinstance(line[1], (list, tuple)) else str(line[1])
                        lines.append(txt)
                raw_text = "\n".join(lines)
    except Exception:
        logger.exception("OCR extraction failed (engine=%s)", name)
        return "", hw_count

    cleaned = _postprocess_ocr_text(raw_text)
    logger.info("OCR complete (%s): %d raw chars → %d cleaned chars",
                name, len(raw_text), len(cleaned))
    return cleaned, hw_count


# --------------------------------------------------------------------------- #
# PDF                                                                         #
# --------------------------------------------------------------------------- #
def _ingest_pdf(path: Path) -> dict:
    """
    Fast hybrid PDF ingestion:
    1. Use PyMuPDF (fitz) for text extraction — 30x faster than pdfplumber
    2. Only OCR pages that have no digital text (scanned/image pages)
    3. Never OCR the entire PDF if most pages are digital
    """
    try:
        import fitz
    except ImportError:
        raise _missing("pymupdf", "PDF", "pip install pymupdf")

    doc = fitz.open(str(path))
    page_count = len(doc)
    text_chunks: list[str] = []
    scanned_page_indices: list[int] = []

    # Phase 1: Extract digital text from all pages (FAST — ~1s for 100+ pages)
    for i, page in enumerate(doc):
        page_text = page.get_text().strip()
        text_chunks.append(page_text)
        # Use multi-signal heuristic: check char count AND content quality
        if len(page_text) < 50:
            scanned_page_indices.append(i)
        elif page_text:
            # Check for garbage text layers (low char variety / low word ratio)
            non_ws = re.sub(r"\s+", "", page_text)
            unique_chars = len(set(non_ws)) if non_ws else 0
            tokens = page_text.split()
            word_like = sum(1 for t in tokens if re.match(r"[a-zA-Z]{2,}", t)) if tokens else 0
            word_ratio = word_like / len(tokens) if tokens else 0
            if unique_chars < 10 or word_ratio < 0.15:
                scanned_page_indices.append(i)

    logger.info("PDF: %d pages, %d digital, %d scanned",
                page_count, page_count - len(scanned_page_indices), len(scanned_page_indices))

    # Phase 2: OCR all scanned pages — never skip OCR for pages that need it
    total_digital_chars = sum(len(text_chunks[i]) for i in range(page_count) if i not in scanned_page_indices)
    skip_ocr = False  # Always OCR pages that are detected as scanned

    if scanned_page_indices and not skip_ocr:
        engine = _get_shared_ocr()
        if engine:
            for idx in scanned_page_indices:
                try:
                    page = doc[idx]
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("png")
                    result = engine.run(img_bytes, filename=f"page_{idx+1}.png")
                    ocr_text = result.get("full_text", "")
                    if ocr_text and len(ocr_text.strip()) > 10:
                        text_chunks[idx] = ocr_text
                        logger.info("  OCR page %d: %d chars", idx+1, len(ocr_text))
                except Exception as e:
                    logger.warning("  OCR page %d failed: %s", idx+1, e)
    elif skip_ocr and scanned_page_indices:
        logger.info("  Skipping OCR on %d pages (enough digital text: %d chars)",
                    len(scanned_page_indices), total_digital_chars)

    doc.close()
    text = "\n".join(text_chunks).strip()

    return {
        "source_type": "pdf",
        "is_scanned": len(scanned_page_indices) > 0,
        "text": text,
        "structured": None,
        "metadata": {
            "pages": page_count,
            "pages_scanned": len(scanned_page_indices),
            "pages_skipped": 0,
            "ocr_used": len(scanned_page_indices) > 0,
            "ocr_engine": "PyMuPDF+EasyOCR" if scanned_page_indices else "PyMuPDF",
        },
    }

    # Fallback: local OCR with pdf2image + poppler
    try:
        from pdf2image import convert_from_path
        import cv2
        import numpy as np
    except ImportError:
        raise IngestionError(
            "This PDF appears to be scanned (no text layer found) and OCR "
            "dependencies are not installed.\n"
            "  Install: pip install pdf2image opencv-python-headless easyocr\n"
            "  Plus Poppler system binary: "
            "https://github.com/oschwartz10612/poppler-windows/releases"
        )

    try:
        kwargs: dict[str, Any] = {"dpi": ocr_config.PDF_DPI}
        if ocr_config.PDF_POPPLER_PATH:
            kwargs["poppler_path"] = ocr_config.PDF_POPPLER_PATH
        images = convert_from_path(str(path), **kwargs)
    except Exception as e:
        raise IngestionError(
            f"pdf2image failed (usually means Poppler isn't installed): {e}\n"
            "  Download Poppler: "
            "https://github.com/oschwartz10612/poppler-windows/releases"
        )

    # MAX_OCR_PAGES: 0 = process all pages, >0 = limit
    max_pages = getattr(ocr_config, "MAX_OCR_PAGES", 0)
    total_pages = len(images)
    pages_to_process = total_pages if max_pages <= 0 else min(total_pages, max_pages)

    logger.info("Scanned PDF: %d page(s) at %d DPI — processing %d pages",
                total_pages, ocr_config.PDF_DPI, pages_to_process)

    ocr_chunks = []
    total_hw_regions = 0
    scanned_pages = 0
    for i, pil_img in enumerate(images[:pages_to_process]):
        logger.info("OCR page %d/%d ...", i + 1, pages_to_process)
        cv_img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
        page_text, hw_count = _ocr_image(cv_img)
        total_hw_regions += hw_count
        scanned_pages += 1
        if page_text:
            ocr_chunks.append(page_text)
        else:
            logger.warning("Page %d returned empty text", i + 1)

    skipped = total_pages - pages_to_process
    if skipped > 0:
        ocr_chunks.append(f"\n[... {skipped} additional pages not processed ...]")

    return {
        "source_type": "pdf",
        "is_scanned": True,
        "text": "\n".join(ocr_chunks),
        "structured": tables_data or None,
        "metadata": {
            "pages": page_count,
            "pages_scanned": scanned_pages,
            "pages_skipped": skipped,
            "ocr_used": True,
            "ocr_engine": get_ocr_engine_name(),
            "handwritten_regions": total_hw_regions,
            "tables_extracted": len(tables_data),
        },
    }


# --------------------------------------------------------------------------- #
# Image                                                                       #
# --------------------------------------------------------------------------- #
def _ingest_image(path: Path) -> dict:
    try:
        import cv2
        import numpy as np
    except ImportError:
        raise _missing("opencv-python", "images",
                       "pip install opencv-python-headless")

    img = cv2.imread(str(path))
    if img is None:
        try:
            from PIL import Image
        except ImportError:
            raise IngestionError(
                f"Could not read {path.name} with OpenCV, and Pillow fallback "
                f"is unavailable. Install: pip install Pillow"
            )
        pil = Image.open(path).convert("RGB")
        img = cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)

    text, hw_count = _ocr_image(img)
    return {
        "source_type": "image",
        "is_scanned": True,
        "text": text,
        "structured": None,
        "metadata": {
            "ocr_used": True,
            "ocr_engine": get_ocr_engine_name(),
            "shape": list(img.shape),
            "handwritten_regions": hw_count,
        },
    }


# --------------------------------------------------------------------------- #
# Image from bytes (for API uploads without saving to disk)                   #
# --------------------------------------------------------------------------- #
def _ingest_image_bytes(data: bytes) -> dict:
    """OCR an image passed as raw bytes (e.g. from an UploadFile)."""
    # Try shared OCR engine first
    shared_text = _shared_ocr_extract(data, "image_upload")
    if shared_text:
            return {
                "source_type": "image",
                "is_scanned": True,
                "text": shared_text,
                "structured": None,
                "metadata": {
                    "ocr_used": True,
                    "ocr_engine": "shared_ocr_service",
                    "handwritten_regions": 0,
                },
            }

    try:
        import cv2
        import numpy as np
    except ImportError:
        raise _missing("opencv-python", "images",
                       "pip install opencv-python-headless")

    arr = np.frombuffer(data, dtype=np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if img is None:
        raise IngestionError("Could not decode image bytes")

    text, hw_count = _ocr_image(img)
    return {
        "source_type": "image",
        "is_scanned": True,
        "text": text,
        "structured": None,
        "metadata": {
            "ocr_used": True,
            "ocr_engine": get_ocr_engine_name(),
            "shape": list(img.shape),
            "handwritten_regions": hw_count,
        },
    }


# --------------------------------------------------------------------------- #
# DOCX                                                                        #
# --------------------------------------------------------------------------- #
def _ingest_docx(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise _missing("python-docx", "DOCX", "pip install python-docx")

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return {
        "source_type": "docx",
        "is_scanned": False,
        "text": "\n".join(parts),
        "structured": None,
        "metadata": {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
    }


# --------------------------------------------------------------------------- #
# XLSX                                                                        #
# --------------------------------------------------------------------------- #
def _ingest_xlsx(path: Path) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise _missing("openpyxl", "XLSX", "pip install openpyxl")

    wb = load_workbook(str(path), data_only=True, read_only=True)
    text_parts: list[str] = []
    rows: list[dict] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_rows = list(ws.iter_rows(values_only=True))
        if not sheet_rows:
            continue
        header = [str(c) if c is not None else f"col_{i}"
                  for i, c in enumerate(sheet_rows[0])]
        text_parts.append(f"## Sheet: {sheet_name}")
        text_parts.append(" | ".join(header))
        for r in sheet_rows[1:]:
            values = [str(c) if c is not None else "" for c in r]
            text_parts.append(" | ".join(values))
            rows.append({"_sheet": sheet_name, **dict(zip(header, values))})

    return {
        "source_type": "xlsx",
        "is_scanned": False,
        "text": "\n".join(text_parts),
        "structured": rows,
        "metadata": {"sheets": wb.sheetnames, "rows": len(rows)},
    }


# --------------------------------------------------------------------------- #
# CSV / TSV                                                                   #
# --------------------------------------------------------------------------- #
def _ingest_csv(path: Path) -> dict:
    raw = _read_text_with_fallback(path)
    try:
        dialect = csv.Sniffer().sniff(raw[:4096], delimiters=",;\t|")
    except csv.Error:
        class _D(csv.excel):
            delimiter = ","
        dialect = _D()

    reader = csv.DictReader(io.StringIO(raw), dialect=dialect)
    rows = list(reader)
    text_lines: list[str] = []
    if reader.fieldnames:
        text_lines.append(" | ".join(reader.fieldnames))
    for row in rows:
        text_lines.append(" | ".join(str(v) for v in row.values()))

    return {
        "source_type": "csv",
        "is_scanned": False,
        "text": "\n".join(text_lines),
        "structured": rows,
        "metadata": {"rows": len(rows), "columns": reader.fieldnames or []},
    }


# --------------------------------------------------------------------------- #
# JSON                                                                        #
# --------------------------------------------------------------------------- #
def _ingest_json(path: Path) -> dict:
    raw = _read_text_with_fallback(path)
    data = json.loads(raw)

    lines: list[str] = []

    def walk(obj: Any, prefix: str = "") -> None:
        if isinstance(obj, dict):
            for k, v in obj.items():
                walk(v, f"{prefix}{k}.")
        elif isinstance(obj, list):
            for i, v in enumerate(obj):
                walk(v, f"{prefix}{i}.")
        else:
            lines.append(f"{prefix.rstrip('.')}: {obj}")

    walk(data)
    structured = data if isinstance(data, list) else None
    return {
        "source_type": "json",
        "is_scanned": False,
        "text": "\n".join(lines),
        "structured": structured,
        "metadata": {"top_level_type": type(data).__name__},
    }


# --------------------------------------------------------------------------- #
# HTML                                                                        #
# --------------------------------------------------------------------------- #
_HTML_TAG = re.compile(r"<[^>]+>")
_WS = re.compile(r"\s+")


def _ingest_html(path: Path) -> dict:
    raw = _read_text_with_fallback(path)
    raw = re.sub(r"<script\b[^>]*>.*?</script>", " ", raw,
                 flags=re.DOTALL | re.IGNORECASE)
    raw = re.sub(r"<style\b[^>]*>.*?</style>", " ", raw,
                 flags=re.DOTALL | re.IGNORECASE)
    text = _HTML_TAG.sub(" ", raw)
    text = _WS.sub(" ", text).strip()
    return {
        "source_type": "html",
        "is_scanned": False,
        "text": text,
        "structured": None,
        "metadata": {},
    }


# --------------------------------------------------------------------------- #
# TXT                                                                         #
# --------------------------------------------------------------------------- #
def _ingest_text(path: Path) -> dict:
    text = _read_text_with_fallback(path)
    return {
        "source_type": "text",
        "is_scanned": False,
        "text": text,
        "structured": None,
        "metadata": {"chars": len(text)},
    }


def _read_text_with_fallback(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="ignore")


# --------------------------------------------------------------------------- #
# Public entry points                                                         #
# --------------------------------------------------------------------------- #
SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}

_DISPATCH = {
    ".pdf":  _ingest_pdf,
    ".docx": _ingest_docx,
    ".xlsx": _ingest_xlsx,
    ".xlsm": _ingest_xlsx,
    ".csv":  _ingest_csv,
    ".tsv":  _ingest_csv,
    ".json": _ingest_json,
    ".html": _ingest_html,
    ".htm":  _ingest_html,
    ".txt":  _ingest_text,
    ".md":   _ingest_text,
    ".log":  _ingest_text,
}


def ingest(source: PathLike) -> dict:
    """Dispatch an ingestion based on file extension."""
    p = Path(source)
    if not p.exists():
        raise FileNotFoundError(p)

    ext = p.suffix.lower()
    if ext in SUPPORTED_IMAGE_EXTS:
        result = _ingest_image(p)
    elif ext in _DISPATCH:
        result = _DISPATCH[ext](p)
    else:
        try:
            result = _ingest_text(p)
        except Exception as e:
            raise IngestionError(f"Unsupported file type: {ext}") from e

    result["source"] = str(p.resolve())
    return result


def ingest_bytes(data: bytes, filename: str) -> dict:
    """
    Ingest from raw bytes + filename (for API uploads).
    Images are decoded directly in memory; other types are written
    to a temp file first.
    """
    import shutil
    import tempfile

    ext = Path(filename).suffix.lower()

    # Images can be decoded from bytes without a temp file
    if ext in SUPPORTED_IMAGE_EXTS:
        result = _ingest_image_bytes(data)
        result["source"] = filename
        return result

    # Everything else needs a temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    try:
        result = ingest(tmp_path)
        result["source"] = filename  # replace temp path with original name
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    return result


def ingest_text(text: str, label: str = "<inline>") -> dict:
    """Ingest a raw string."""
    return {
        "source": label,
        "source_type": "text",
        "is_scanned": False,
        "text": text,
        "structured": None,
        "metadata": {"chars": len(text)},
    }


def ingest_as(source: PathLike, forced_type: str) -> dict:
    """Force a particular handler when the extension lies."""
    p = Path(source)
    if not p.exists():
        raise FileNotFoundError(p)

    forced = forced_type.lower().strip().lstrip(".")
    handlers = {
        "pdf": _ingest_pdf,
        "image": _ingest_image,
        "docx": _ingest_docx,
        "xlsx": _ingest_xlsx,
        "csv": _ingest_csv,
        "tsv": _ingest_csv,
        "json": _ingest_json,
        "html": _ingest_html,
        "txt": _ingest_text,
        "text": _ingest_text,
    }
    if forced not in handlers:
        raise ValueError(f"Unknown forced type: {forced_type}")
    result = handlers[forced](p)
    result["source"] = str(p.resolve())
    return result


def preload_ocr() -> None:
    """
    Pre-load the OCR model at server startup so the first request
    doesn't pay the cold-start penalty. Safe to call multiple times.
    """
    try:
        _get_ocr_reader()
        logger.info("OCR engine pre-loaded: %s", get_ocr_engine_name())
    except Exception:
        logger.warning("OCR pre-load failed — will retry on first image request")


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s | %(name)s | %(message)s")
    if len(sys.argv) < 2:
        print("Usage: python ingestion.py <file>")
        sys.exit(1)
    out = ingest(sys.argv[1])
    print(f"[{out['source_type']}] scanned={out['is_scanned']} "
          f"chars={len(out['text'])}")
    if out["metadata"].get("ocr_engine"):
        print(f"OCR engine: {out['metadata']['ocr_engine']}")
    print("---")
    print(out["text"][:2000])
